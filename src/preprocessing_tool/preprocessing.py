import argparse
import os
import pandas as pd
from docx import Document
from transformers import pipeline

# Load Hugging Face question-answering model
qa_pipeline = pipeline("question-answering", model="deepset/roberta-base-squad2")

# Define the functions as before (loading questionnaire, transcript, extracting answers, generating Excel)
def load_questionnaire(excel_file):
    df = pd.read_excel(excel_file, header=None)
    questions = df.iloc[:, 0].tolist()
    return questions

def load_transcript(file_path):
    """Extracts text from paragraphs and tables in a .docx file while maintaining order."""
    doc = Document(file_path)
    extracted_text = []

    # Iterate through all elements in the document body while maintaining order
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraphs
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if para and para.text.strip():
                extracted_text.append(("para", para.text.strip()))  # Add a tag for paragraph

        elif element.tag.endswith('tbl'):  # Tables
            table = next((t for t in doc.tables if t._element == element), None)
            if table:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        extracted_text.append(("table", " | ".join(row_text)))  # Add a tag for table row

    return extracted_text

def extract_answers(questions, transcript):
    answers = {}
    for question in questions:
        try:
            result = qa_pipeline(question=question, context=transcript)
            answers[question] = result['answer'] if result['score'] > 0.5 else "Uncertain / Not Found"
        except Exception as e:
            answers[question] = "Error: " + str(e)
    return answers


def get_answer_from_context(question_match, doc):
    """Extract answer by getting the next sentences after the question."""
    answer = []
    # Find the index of the question sentence
    question_index = list(doc.sents).index(question_match)
    
    # Consider the next few sentences (you could adjust this based on your data)
    for sent in list(doc.sents)[question_index+1:question_index+3]:  # Get next 2 sentences after the question
        answer.append(sent.text.strip())
    
    return " ".join(answer)

def generate_excel_output(answers, output_file):
    df = pd.DataFrame(list(answers.items()), columns=['Question', 'Answer'])
    df.to_excel(output_file, index=False)

# Main function to handle CLI input/output
def main():
    # Set up CLI argument parsing
    parser = argparse.ArgumentParser(description="Process interview transcript based on a questionnaire.")
    parser.add_argument('questionnaire_path', type=str, help="Path to the questionnaire Excel file")
    parser.add_argument('transcript_path', type=str, help="Path to the interview transcript Word file")
    
    args = parser.parse_args()
    
    # Check if input files exist
    if not os.path.exists(args.questionnaire_path):
        print(f"Error: The questionnaire file '{args.questionnaire_path}' does not exist.")
        return
    if not os.path.exists(args.transcript_path):
        print(f"Error: The transcript file '{args.transcript_path}' does not exist.")
        return
    
    # Step 1: Load the questionnaire
    questions = load_questionnaire(args.questionnaire_path)
    print("################### LOADING TRANSCRIPT ###################\n")
    
    # Step 2: Load the interview transcript
    transcript = load_transcript(args.transcript_path)
    
    # Step 3: Extract answers
    answers = extract_answers(questions, transcript)
    
    # Step 4: Prepare the output folder
    output_dir = os.path.join('data', 'processed_data')
    os.makedirs(output_dir, exist_ok=True)
    output_file = os.path.join(output_dir, 'processed_interviews.xlsx')
    
    # Step 5: Generate output Excel file
    generate_excel_output(answers, output_file)
    print(f"Output saved to {output_file}")

if __name__ == "__main__":
    main()
